"""Generate Documentation.docx with all model figures embedded."""

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
import joblib
import shap
import os
from sklearn.linear_model import LinearRegression
from sklearn.ensemble import RandomForestRegressor
from sklearn.preprocessing import LabelEncoder
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
from sklearn.inspection import PartialDependenceDisplay
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sns.set_theme(style='whitegrid', palette='muted')
IMG_DIR = 'doc_images'
os.makedirs(IMG_DIR, exist_ok=True)

# ── Load data and model ──────────────────────────────────────────────────────
print("Loading data and model...")
df = pd.read_csv('SriLanka_Weather_Dataset.csv')
artifacts = joblib.load('model_artifacts.joblib')
xgb_model = artifacts['model']
feature_cols = artifacts['feature_cols']
le_city = artifacts['label_encoder_city']

df['time_dt'] = pd.to_datetime(df['time'])
df['month'] = df['time_dt'].dt.month
df['day_of_year'] = df['time_dt'].dt.dayofyear
df['year'] = df['time_dt'].dt.year
df['sunrise_dt'] = pd.to_datetime(df['sunrise'])
df['sunset_dt'] = pd.to_datetime(df['sunset'])
df['day_length_hours'] = (df['sunset_dt'] - df['sunrise_dt']).dt.total_seconds() / 3600
wind_dir_rad = np.deg2rad(df['winddirection_10m_dominant'])
df['wind_dir_sin'] = np.sin(wind_dir_rad)
df['wind_dir_cos'] = np.cos(wind_dir_rad)
df['city_encoded'] = le_city.transform(df['city'])

X = df[feature_cols]
y = df['precipitation_sum']

sorted_idx = df['time_dt'].argsort().values
X_sorted = X.iloc[sorted_idx].reset_index(drop=True)
y_sorted = y.iloc[sorted_idx].reset_index(drop=True)
n = len(X_sorted)
train_end = int(n * 0.8)
val_end = int(n * 0.9)
X_train, y_train = X_sorted.iloc[:train_end], y_sorted.iloc[:train_end]
X_test, y_test = X_sorted.iloc[val_end:], y_sorted.iloc[val_end:]

# Baselines
print("Training baselines...")
lr = LinearRegression().fit(X_train, y_train)
lr_pred = lr.predict(X_test)
rf = RandomForestRegressor(n_estimators=200, max_depth=15, random_state=42, n_jobs=-1)
rf.fit(X_train, y_train)
rf_pred = rf.predict(X_test)
xgb_pred = xgb_model.predict(X_test)

models_data = [
    ('Linear Regression', lr_pred),
    ('Random Forest', rf_pred),
    ('XGBoost (Tuned)', xgb_pred),
]

# ── Generate figures ─────────────────────────────────────────────────────────

def savefig(name):
    path = os.path.join(IMG_DIR, name)
    plt.savefig(path, dpi=150, bbox_inches='tight')
    plt.close()
    return path

# Fig 1: Precipitation distribution
print("Generating figures...")
fig, axes = plt.subplots(1, 3, figsize=(18, 5))
axes[0].hist(df['precipitation_sum'], bins=80, edgecolor='black', alpha=0.7, color='steelblue')
axes[0].set_title('Distribution of Daily Precipitation')
axes[0].set_xlabel('Precipitation (mm)')
axes[0].set_ylabel('Frequency')
axes[0].axvline(df['precipitation_sum'].mean(), color='red', linestyle='--',
                label=f"Mean = {df['precipitation_sum'].mean():.1f} mm")
axes[0].legend()

monthly_avg = df.groupby(df['time_dt'].dt.month)['precipitation_sum'].mean()
axes[1].bar(monthly_avg.index, monthly_avg.values, color='teal', edgecolor='black', alpha=0.8)
axes[1].set_title('Mean Precipitation by Month')
axes[1].set_xlabel('Month')
axes[1].set_ylabel('Mean Precipitation (mm)')
axes[1].set_xticks(range(1, 13))
axes[1].set_xticklabels(['Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec'], rotation=45)

city_avg = df.groupby('city')['precipitation_sum'].mean().sort_values(ascending=True)
axes[2].barh(city_avg.index, city_avg.values, color='coral', edgecolor='black', alpha=0.8)
axes[2].set_title('Mean Precipitation by City')
axes[2].set_xlabel('Mean Precipitation (mm)')
plt.tight_layout()
fig1 = savefig('fig1_eda.png')

# Fig 2: Correlation with target
numeric_features = X.select_dtypes(include=[np.number]).columns.tolist()
corr = X[numeric_features].corrwith(y).sort_values(ascending=False)
fig, ax = plt.subplots(figsize=(10, 8))
colors = ['steelblue' if v >= 0 else 'coral' for v in corr.values]
ax.barh(corr.index, corr.values, color=colors, edgecolor='black', alpha=0.8)
ax.set_title('Pearson Correlation of Features with Precipitation', fontsize=14)
ax.set_xlabel('Correlation Coefficient')
ax.axvline(0, color='black', linewidth=0.8)
plt.tight_layout()
fig2 = savefig('fig2_correlation.png')

# Fig 3: Actual vs Predicted
fig, axes = plt.subplots(1, 3, figsize=(20, 6))
for ax, (name, pred) in zip(axes, models_data):
    ax.scatter(y_test, pred, alpha=0.15, s=5, color='steelblue')
    lims = [0, max(y_test.max(), max(pred))]
    ax.plot(lims, lims, 'r--', linewidth=1.5, label='Perfect prediction')
    ax.set_title(f'{name}\nR² = {r2_score(y_test, pred):.4f}')
    ax.set_xlabel('Actual Precipitation (mm)')
    ax.set_ylabel('Predicted Precipitation (mm)')
    ax.legend()
plt.suptitle('Actual vs. Predicted Precipitation — Model Comparison', fontsize=14, y=1.02)
plt.tight_layout()
fig3 = savefig('fig3_actual_vs_pred.png')

# Fig 4: Residual distributions
fig, axes = plt.subplots(1, 3, figsize=(20, 5))
for ax, (name, pred) in zip(axes, models_data):
    residuals = y_test.values - pred
    ax.hist(residuals, bins=80, edgecolor='black', alpha=0.7, color='mediumpurple')
    ax.axvline(0, color='red', linestyle='--')
    ax.set_title(f'{name}\nMean residual = {residuals.mean():.3f}')
    ax.set_xlabel('Residual (mm)')
    ax.set_ylabel('Frequency')
plt.suptitle('Residual Distributions', fontsize=14, y=1.02)
plt.tight_layout()
fig4 = savefig('fig4_residuals.png')

# Fig 5: Metric comparison bars
results = {
    'Model': ['Linear Regression', 'Random Forest', 'XGBoost (Tuned)'],
    'RMSE': [np.sqrt(mean_squared_error(y_test, p)) for _, p in models_data],
    'MAE': [mean_absolute_error(y_test, p) for _, p in models_data],
    'R²': [r2_score(y_test, p) for _, p in models_data],
}
results_df = pd.DataFrame(results).set_index('Model')
fig, axes = plt.subplots(1, 3, figsize=(16, 5))
for ax, metric in zip(axes, ['RMSE', 'MAE', 'R²']):
    values = results_df[metric].values
    colors = ['#e74c3c', '#f39c12', '#2ecc71']
    bars = ax.bar(results_df.index, values, color=colors, edgecolor='black', alpha=0.85)
    ax.set_title(metric, fontsize=14, fontweight='bold')
    ax.set_ylabel(metric)
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01*max(values),
                f'{val:.3f}', ha='center', va='bottom', fontweight='bold')
    ax.tick_params(axis='x', rotation=15)
plt.suptitle('Performance Metric Comparison', fontsize=14, y=1.02)
plt.tight_layout()
fig5 = savefig('fig5_metrics.png')

# Fig 6: Feature importance
importance = xgb_model.feature_importances_
feat_imp = pd.Series(importance, index=feature_cols).sort_values(ascending=True)
fig, ax = plt.subplots(figsize=(10, 8))
feat_imp.plot(kind='barh', ax=ax, color='teal', edgecolor='black', alpha=0.8)
ax.set_title('XGBoost Built-in Feature Importance (Gain)', fontsize=14)
ax.set_xlabel('Importance (Gain)')
plt.tight_layout()
fig6 = savefig('fig6_feature_importance.png')

# Fig 7: SHAP bar + beeswarm
print("Computing SHAP values (this may take a minute)...")
explainer = shap.TreeExplainer(xgb_model)
X_test_sample = X_test.sample(n=min(2000, len(X_test)), random_state=42)
shap_values = explainer.shap_values(X_test_sample)

fig, ax = plt.subplots(figsize=(10, 8))
shap.summary_plot(shap_values, X_test_sample, plot_type='bar', show=False)
plt.title('SHAP Global Feature Importance', fontsize=14)
plt.tight_layout()
fig7a = savefig('fig7a_shap_bar.png')

fig, ax = plt.subplots(figsize=(10, 8))
shap.summary_plot(shap_values, X_test_sample, show=False)
plt.title('SHAP Beeswarm Plot', fontsize=14)
plt.tight_layout()
fig7b = savefig('fig7b_shap_beeswarm.png')

# Fig 8: SHAP dependence plots
top_features = pd.Series(np.abs(shap_values).mean(axis=0), index=feature_cols).sort_values(ascending=False)
top_4 = top_features.head(4).index.tolist()

fig, axes = plt.subplots(2, 2, figsize=(16, 12))
for ax, feat in zip(axes.flatten(), top_4):
    shap.dependence_plot(feat, shap_values, X_test_sample, ax=ax, show=False)
    ax.set_title(f'SHAP Dependence: {feat}', fontsize=12)
plt.suptitle('SHAP Dependence Plots — Top 4 Features', fontsize=14, y=1.01)
plt.tight_layout()
fig8 = savefig('fig8_shap_dependence.png')

# Fig 9: Partial Dependence Plots
print("Generating PDP plots...")
top_4_indices = [feature_cols.index(f) for f in top_4]
fig, axes = plt.subplots(2, 2, figsize=(16, 10))
X_train_sample = X_train.sample(n=5000, random_state=42)
for ax, feat_idx, feat_name in zip(axes.flatten(), top_4_indices, top_4):
    PartialDependenceDisplay.from_estimator(
        xgb_model, X_train_sample, features=[feat_idx],
        feature_names=feature_cols, ax=ax, kind='average'
    )
    ax.set_title(f'PDP: {feat_name}', fontsize=12)
plt.suptitle('Partial Dependence Plots — Top 4 Features', fontsize=14, y=1.01)
plt.tight_layout()
fig9 = savefig('fig9_pdp.png')

# Fig 10: SHAP waterfall for one high-rain sample
high_rain_idx = X_test_sample.index[y_test.loc[X_test_sample.index].argmax()]
sample_pos = list(X_test_sample.index).index(high_rain_idx)
sv = shap_values[sample_pos]
shap_s = pd.Series(sv, index=feature_cols)
top_shap = shap_s.reindex(shap_s.abs().sort_values(ascending=False).head(15).index).sort_values()

fig, ax = plt.subplots(figsize=(10, 7))
colors = ["#ff0051" if v > 0 else "#008bfb" for v in top_shap.values]
ax.barh(range(len(top_shap)), top_shap.values, color=colors, edgecolor="none", height=0.6)
for i, (feat, val) in enumerate(zip(top_shap.index, top_shap.values)):
    input_val = X_test_sample.loc[high_rain_idx, feat]
    label = f"{feat} = {input_val:.2f}"
    ax.text(-0.01 if val > 0 else 0.01, i, label,
            ha="right" if val > 0 else "left", va="center", fontsize=9)
ax.set_yticks([])
ax.set_xlabel("SHAP value (impact on prediction)")
actual_val = y_test.loc[high_rain_idx]
pred_val = xgb_model.predict(X_test_sample.iloc[[sample_pos]])[0]
ax.set_title(f'SHAP Local Explanation — Actual: {actual_val:.1f} mm, Predicted: {pred_val:.1f} mm',
             fontsize=12, fontweight='bold')
ax.axvline(0, color='black', linewidth=0.8)
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
red_p = plt.Line2D([0], [0], color="#ff0051", lw=6, label="Pushes UP (more rain)")
blue_p = plt.Line2D([0], [0], color="#008bfb", lw=6, label="Pushes DOWN (less rain)")
ax.legend(handles=[red_p, blue_p], loc="lower right", fontsize=9)
plt.tight_layout()
fig10 = savefig('fig10_shap_waterfall.png')

print("All figures generated.\n")

# ── Build DOCX ───────────────────────────────────────────────────────────────
print("Building Documentation.docx...")
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)

def add_figure(img_path, caption, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()

# ── Title page ──
title = doc.add_heading('Predicting Daily Precipitation Levels in Sri Lankan Cities\nfor Flood Risk Assessment', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Machine Learning Assignment').bold = True
doc.add_paragraph()
doc.add_page_break()

# ── Table of Contents ──
add_heading('Table of Contents', 1)
toc_items = [
    '1. Problem Description',
    '2. Methodology',
    '   2.1 Dataset Description',
    '   2.2 Feature Selection',
    '   2.3 Data Preprocessing',
    '   2.4 Algorithm Selection: XGBoost',
    '   2.5 Model Training Strategy',
    '   2.6 Hyperparameter Tuning',
    '3. Results',
    '   3.1 Performance Metrics',
    '   3.2 Model Comparison',
    '   3.3 Visual Analysis',
    '4. Interpretation',
    '   4.1 Feature Importance',
    '   4.2 SHAP Analysis',
    '   4.3 Partial Dependence Plots',
    '   4.4 Domain Alignment',
    '5. Discussion',
    '   5.1 Model Limitations',
    '   5.2 Data Quality Concerns',
    '   5.3 Bias and Fairness',
    '   5.4 Ethical Implications',
    '   5.5 Future Work',
    '6. Front-End Integration',
    '7. References',
]
for item in toc_items:
    doc.add_paragraph(item)
doc.add_page_break()

# ── 1. Problem Description ──
add_heading('1. Problem Description', 1)

add_heading('1.1 Background', 2)
add_para(
    'Flooding is one of the most frequent and destructive natural disasters in Sri Lanka. '
    'Driven by intense monsoon rainfall and tropical weather systems, flood events routinely '
    'cause significant damage to infrastructure, agriculture, and human life. The country '
    'experiences two major monsoon seasons \u2014 the Southwest monsoon (May\u2013September) and '
    'the Northeast monsoon (October\u2013January) \u2014 both of which bring sustained periods '
    'of heavy rainfall to different geographic regions of the island.'
)
add_para(
    'Accurate prediction of daily precipitation levels is a critical prerequisite for effective '
    'flood risk assessment. With reliable precipitation forecasts, disaster management authorities '
    'can issue timely warnings, allocate emergency resources, and minimize the impact of flood '
    'events on vulnerable communities.'
)

add_heading('1.2 Objective', 2)
add_para(
    'The objective of this project is to build a machine learning model that predicts daily '
    'total precipitation (in millimeters) across 30 Sri Lankan cities using meteorological, '
    'geographic, and temporal features. The predicted precipitation values can be classified '
    'into flood risk categories (Low, Moderate, High, Very High) to support practical decision-making.'
)

add_heading('1.3 Real-World Relevance', 2)
add_bullet('Early warning systems: High precipitation predictions can trigger alerts for flood-prone regions.')
add_bullet('Disaster resource allocation: Authorities can pre-position emergency supplies and personnel based on predicted rainfall intensity.')
add_bullet('Agricultural planning: Farmers can make informed decisions about planting, irrigation, and harvest timing.')
add_bullet('Urban infrastructure: City planners can assess drainage system adequacy and plan maintenance schedules around high-rainfall periods.')

doc.add_page_break()

# ── 2. Methodology ──
add_heading('2. Methodology', 1)

add_heading('2.1 Dataset Description', 2)
add_table(
    ['Property', 'Details'],
    [
        ['Data Source', 'Open-Meteo Historical Weather API (publicly available, open-source)'],
        ['Geographic Coverage', '30 cities across Sri Lanka'],
        ['Temporal Coverage', 'January 1, 2010 \u2014 June 17, 2023'],
        ['Total Records', '147,480 daily observations'],
        ['Original Columns', '24'],
        ['Ethical Considerations', 'No personal or sensitive data; entirely public weather records'],
    ]
)
add_para(
    'The dataset contains daily weather observations for 30 Sri Lankan cities including '
    'Colombo, Kandy, Galle, Jaffna, Trincomalee, Ratnapura, Badulla, and others spanning '
    'both the wet zone (southwestern Sri Lanka) and the dry zone (northern and eastern regions). '
    'Each record includes temperature readings, apparent temperature, solar radiation, wind '
    'measurements, evapotranspiration, precipitation totals, and geographic metadata.'
)

add_heading('Exploratory Data Analysis', 3)
add_figure(fig1, 'Figure 1: Precipitation distribution, monthly averages, and per-city averages.')
add_figure(fig2, 'Figure 2: Pearson correlation of features with precipitation.', width=5.5)

add_heading('2.2 Feature Selection', 2)

add_para('Dependent Variable (Target)', bold=True)
add_para(
    'precipitation_sum \u2014 Total daily precipitation in millimeters (continuous variable, regression target). '
    'Key statistics: Mean = 5.98 mm, Maximum = 338.8 mm, Zero-precipitation days = 18.1% of all records. '
    'The distribution is right-skewed, with the majority of days having low to moderate precipitation.'
)

add_para('Independent Variables (20 Features)', bold=True)
add_table(
    ['Category', 'Feature', 'Description'],
    [
        ['Meteorological', 'temperature_2m_max', 'Daily maximum temperature at 2m (\u00b0C)'],
        ['', 'temperature_2m_min', 'Daily minimum temperature at 2m (\u00b0C)'],
        ['', 'temperature_2m_mean', 'Daily mean temperature at 2m (\u00b0C)'],
        ['', 'apparent_temperature_max', 'Maximum "feels-like" temperature (\u00b0C) \u2014 encodes humidity'],
        ['', 'apparent_temperature_min', 'Minimum "feels-like" temperature (\u00b0C)'],
        ['', 'apparent_temperature_mean', 'Mean "feels-like" temperature (\u00b0C)'],
        ['', 'shortwave_radiation_sum', 'Total shortwave solar radiation (MJ/m\u00b2)'],
        ['', 'windspeed_10m_max', 'Maximum wind speed at 10m height (km/h)'],
        ['', 'windgusts_10m_max', 'Maximum wind gusts at 10m height (km/h)'],
        ['', 'et0_fao_evapotranspiration', 'FAO reference evapotranspiration (mm)'],
        ['Geographic', 'latitude', 'City latitude'],
        ['', 'longitude', 'City longitude'],
        ['', 'elevation', 'City elevation above sea level (m)'],
        ['', 'city_encoded', 'Label-encoded city identifier'],
        ['Temporal', 'month', 'Month of year (1\u201312)'],
        ['', 'day_of_year', 'Day of year (1\u2013366)'],
        ['', 'year', 'Calendar year'],
        ['', 'day_length_hours', 'Daylight duration from sunrise/sunset'],
        ['Engineered', 'wind_dir_sin', 'Sine of wind direction (circular encoding)'],
        ['', 'wind_dir_cos', 'Cosine of wind direction (circular encoding)'],
    ]
)

add_para('Excluded Variables and Rationale', bold=True)
add_table(
    ['Feature', 'Reason for Exclusion'],
    [
        ['rain_sum', 'Identical to precipitation_sum (snowfall = 0) \u2014 direct data leakage'],
        ['snowfall_sum', 'Constant value of 0 \u2014 zero variance, no predictive information'],
        ['precipitation_hours', 'Hours of precipitation on the same day \u2014 direct data leakage'],
        ['weathercode', 'Encodes observed weather condition \u2014 data leakage'],
        ['sunrise / sunset', 'Replaced by derived day_length_hours'],
        ['country', 'Always "Sri Lanka" \u2014 zero variance'],
        ['time', 'Replaced by month, day_of_year, year'],
        ['winddirection_10m_dominant', 'Replaced by wind_dir_sin and wind_dir_cos'],
    ]
)

add_heading('2.3 Data Preprocessing', 2)
add_para('The following preprocessing steps were applied:')
add_bullet('Date parsing and temporal feature engineering: Extracted month, day_of_year, and year from the time column.')
add_bullet('Day length calculation: Computed day_length_hours from sunrise and sunset timestamps.')
add_bullet('Circular encoding of wind direction: Encoded as sin(direction) and cos(direction) to preserve circularity.')
add_bullet('Label encoding of city: Transformed the categorical city column to integer codes.')
add_bullet('Missing value check: No missing values found \u2014 no imputation required.')
add_bullet('Feature scaling: Not applied. XGBoost is tree-based and invariant to feature scaling.')

add_heading('2.4 Algorithm Selection: XGBoost', 2)
add_para(
    'XGBoost (eXtreme Gradient Boosting) was selected as the primary algorithm. '
    'It is a scalable, optimized implementation of gradient-boosted decision trees that builds '
    'an ensemble of weak learners sequentially, where each new tree corrects residual errors.'
)
add_para('Key Characteristics:', bold=True)
add_bullet('Gradient boosting framework: Minimizes squared error by iteratively fitting trees to the negative gradient.')
add_bullet('Built-in regularization: L1 and L2 penalties on leaf weights prevent overfitting.')
add_bullet('Column and row subsampling: Reduces variance and improves generalization.')
add_bullet('Native handling of missing values: Learns optimal split directions automatically.')
add_bullet('Computational efficiency: Parallel tree construction for scalability.')

add_para('Comparison with Standard Algorithms', bold=True)
add_table(
    ['Aspect', 'Linear Regression', 'Decision Tree', 'k-NN', 'XGBoost'],
    [
        ['Model type', 'Single linear model', 'Single tree', 'Instance-based', 'Ensemble of boosted trees'],
        ['Non-linearity', 'Cannot capture', 'Captures but overfits', 'Local averaging', 'Captures via additive trees'],
        ['Feature interactions', 'Manual', 'Via splits', 'Via distance', 'Through tree depth'],
        ['Regularization', 'Ridge/Lasso', 'Pruning', 'k selection', 'L1+L2+complexity'],
        ['Skewed targets', 'Assumes normality', 'Moderate', 'Moderate', 'Robust (non-parametric)'],
        ['Scalability', 'Good', 'Good', 'Poor at scale', 'Excellent'],
    ]
)

add_para('Justification for This Task:', bold=True)
add_bullet('Non-linear relationships: Precipitation depends on complex interactions that linear models cannot capture.')
add_bullet('Skewed target distribution: 18.1% zero-precipitation days, extreme events >300 mm.')
add_bullet('Mixed feature types: Continuous, cyclical, and categorical features handled naturally.')
add_bullet('Regularization: Prevents overfitting with 147K records and 20 features.')
add_bullet('Explainability: Seamless integration with SHAP for global and local explanations.')

add_heading('2.5 Model Training Strategy', 2)
add_para('A temporal split was used to prevent future data leaking into training:')
add_table(
    ['Split', 'Records', 'Percentage', 'Approximate Period'],
    [
        ['Training', '117,984', '80%', '2010 \u2013 mid-2020'],
        ['Validation', '14,748', '10%', 'mid-2020 \u2013 mid-2021'],
        ['Test', '14,748', '10%', 'mid-2021 \u2013 June 2023'],
    ]
)
add_para('Two baseline models were trained: Linear Regression and Random Forest (200 trees, max depth 15).')

add_heading('2.6 Hyperparameter Tuning', 2)
add_para('Optuna (Bayesian optimization with TPE sampler) was used for 50 trials, optimizing validation RMSE.')
add_table(
    ['Hyperparameter', 'Search Range', 'Best Value'],
    [
        ['n_estimators', '200 \u2013 1,000', '900'],
        ['max_depth', '3 \u2013 10', '8'],
        ['learning_rate', '0.01 \u2013 0.3', '0.021'],
        ['subsample', '0.6 \u2013 1.0', '0.695'],
        ['colsample_bytree', '0.5 \u2013 1.0', '0.700'],
        ['reg_alpha (L1)', '0.001 \u2013 10.0', '0.001'],
        ['reg_lambda (L2)', '0.001 \u2013 10.0', '0.278'],
        ['min_child_weight', '1 \u2013 10', '9'],
    ]
)

doc.add_page_break()

# ── 3. Results ──
add_heading('3. Results', 1)

add_heading('3.1 Performance Metrics', 2)
add_para('All metrics are reported on the held-out test set (14,748 samples, mid-2021 to June 2023).')
add_table(
    ['Model', 'RMSE (mm)', 'MAE (mm)', 'R\u00b2'],
    [
        ['Linear Regression', '6.464', '4.166', '0.4272'],
        ['Random Forest', '6.544', '3.251', '0.4130'],
        ['XGBoost (Tuned)', '6.354', '3.123', '0.4465'],
    ]
)

add_heading('3.2 Model Comparison', 2)
add_para(
    'XGBoost (Tuned) achieves the best performance across all three metrics. '
    'Compared to Linear Regression, XGBoost reduces RMSE by 1.7%, MAE by 25.0%, and improves '
    'R\u00b2 by 4.5%. Compared to Random Forest, XGBoost reduces RMSE by 2.9%, MAE by 3.9%, '
    'and improves R\u00b2 by 8.1%.'
)
add_para(
    'The R\u00b2 value of 0.4465 indicates that the model explains approximately 44.7% of '
    'the variance in daily precipitation. This is a reasonable result given the inherent '
    'stochasticity of daily precipitation, which is influenced by micro-scale atmospheric '
    'dynamics that surface-level measurements cannot fully capture.'
)

add_heading('3.3 Visual Analysis', 2)
add_figure(fig3, 'Figure 3: Actual vs. Predicted precipitation scatter plots for all three models.')
add_figure(fig4, 'Figure 4: Residual distribution histograms for all three models.')
add_figure(fig5, 'Figure 5: Performance metric comparison (RMSE, MAE, R\u00b2) across models.')

doc.add_page_break()

# ── 4. Interpretation ──
add_heading('4. Interpretation', 1)

add_heading('4.1 Feature Importance (XGBoost Gain)', 2)
add_table(
    ['Rank', 'Feature', 'Importance'],
    [
        ['1', 'et0_fao_evapotranspiration', '0.1791'],
        ['2', 'shortwave_radiation_sum', '0.1116'],
        ['3', 'temperature_2m_mean', '0.0675'],
        ['4', 'day_length_hours', '0.0533'],
        ['5', 'apparent_temperature_max', '0.0520'],
        ['6', 'apparent_temperature_min', '0.0505'],
        ['7', 'day_of_year', '0.0493'],
        ['8', 'month', '0.0449'],
        ['9', 'windgusts_10m_max', '0.0434'],
        ['10', 'windspeed_10m_max', '0.0414'],
    ]
)
add_figure(fig6, 'Figure 6: XGBoost built-in feature importance ranked by gain.', width=5.5)

add_heading('4.2 SHAP Analysis', 2)
add_para(
    'SHAP (SHapley Additive exPlanations) provides both global (which features matter overall) '
    'and local (why a specific prediction was made) interpretability grounded in game theory.'
)
add_figure(fig7a, 'Figure 7: SHAP global feature importance (mean |SHAP value|).', width=5.5)
add_figure(fig7b, 'Figure 8: SHAP beeswarm plot showing feature value impacts on predictions.', width=5.5)

add_para(
    'Key observations from the SHAP beeswarm plot: Low evapotranspiration (blue) strongly '
    'increases predicted precipitation, consistent with saturated atmospheric conditions. '
    'Low solar radiation increases predicted rain (more cloud cover). Higher apparent '
    'temperatures (encoding humidity) tend to increase predicted precipitation.'
)

add_figure(fig8, 'Figure 9: SHAP dependence plots for the top 4 features.')

add_para('SHAP Local Explanation (Single Prediction)', bold=True)
add_para(
    'The following figure shows a SHAP explanation for a single high-rainfall day, '
    'demonstrating how each feature pushed the prediction from the baseline average '
    'to the final predicted value.'
)
add_figure(fig10, 'Figure 10: SHAP local explanation for a high-rainfall prediction.', width=5.5)

add_heading('4.3 Partial Dependence Plots', 2)
add_para(
    'Partial Dependence Plots show the marginal effect of a feature on the predicted outcome, '
    'averaged over all other features.'
)
add_figure(fig9, 'Figure 11: Partial Dependence Plots for the top 4 features.')
add_para(
    'Key findings: Precipitation decreases sharply as evapotranspiration increases. '
    'Solar radiation shows a clear negative relationship. Mean temperature has a moderate '
    'positive relationship. Day length captures subtle seasonal signals.'
)

add_heading('4.4 Domain Alignment', 2)
add_bullet('Evapotranspiration as top predictor: Low ET indicates high humidity and saturated conditions \u2014 prerequisites for precipitation.')
add_bullet('Solar radiation inverse relationship: Cloudy skies block radiation and are associated with rainfall.')
add_bullet('Seasonal features (month, day_of_year) capture the bimodal monsoon pattern without explicit programming.')
add_bullet('Geographic features differentiate between the wet zone (SW Sri Lanka, 2,500+ mm/yr) and dry zone (N/E, <1,750 mm/yr).')
add_bullet('Wind patterns reflect moisture-laden air masses from the Indian Ocean during monsoon periods.')

doc.add_page_break()

# ── 5. Discussion ──
add_heading('5. Discussion', 1)

add_heading('5.1 Model Limitations', 2)
add_para(
    'Same-day features, not true forecasting: The model uses same-day meteorological '
    'measurements. In a real operational setting, these would not be available in advance. '
    'A true forecasting model would require lagged features or NWP outputs.'
)
add_para(
    'No temporal autocorrelation modeling: Each day is treated independently. Weather '
    'patterns often persist over multiple days. Lagged features or time-series approaches '
    'could capture these dependencies.'
)
add_para(
    'Extreme event underestimation: The model may underpredict rare extreme events (>100 mm) '
    'that are critical for flood risk assessment.'
)
add_para(
    'R\u00b2 of 0.45: 55% of precipitation variance remains unexplained due to micro-scale '
    'atmospheric dynamics that surface measurements cannot capture.'
)

add_heading('5.2 Data Quality Concerns', 2)
add_para('Single data source: All data comes from ERA5 reanalysis via Open-Meteo, which may differ from ground-station measurements.')
add_para('Urban-centric sampling: 30 cities are primarily urban centers; rural and mountainous areas are underrepresented.')
add_para('Temporal boundary: Dataset ends June 2023. Climate change may alter future patterns.')

add_heading('5.3 Bias and Fairness', 2)
add_para('Geographic prediction bias: Accuracy likely varies across cities with different microclimates.')
add_para('No socioeconomic context: The model predicts precipitation, not flood impact. Vulnerable communities may be disproportionately affected.')

add_heading('5.4 Ethical Implications', 2)
add_para('False sense of security: Under-predictions for extreme events could lead to inadequate precautions. Prediction uncertainty must be communicated.')
add_para('Resource allocation risks: Biased predictions could lead to inequitable distribution of disaster response resources.')
add_para('Data transparency: The dataset uses entirely public weather records with no personal information.')

add_heading('5.5 Future Work', 2)
add_bullet('Lagged features: Incorporate previous-day and multi-day rolling averages for true forecasting capability.')
add_bullet('Probabilistic predictions: Use quantile regression or conformal prediction to communicate uncertainty.')
add_bullet('Extreme value modeling: Apply specialized techniques for predicting tail events.')
add_bullet('Spatial features: Incorporate terrain data and upstream catchment information.')
add_bullet('Model retraining pipeline: Automate retraining as new data becomes available.')

doc.add_page_break()

# ── 6. Front-End Integration ──
add_heading('6. Front-End Integration', 1)
add_para(
    'A Streamlit web application (app.py) was developed to make the trained model accessible '
    'to non-technical users. The application provides:'
)
add_bullet('City selection: Dropdown with 30 Sri Lankan cities, auto-populating geographic features.')
add_bullet('Date picker: Automatically derives temporal features (month, day of year, year).')
add_bullet('Weather parameter inputs: Sliders for temperature, radiation, wind, evapotranspiration.')
add_bullet('Precipitation prediction: Displays predicted daily precipitation in millimeters.')
add_bullet('Flood risk indicator: Color-coded risk level (Low <2.5 mm, Moderate 2.5\u201315 mm, High 15\u201350 mm, Very High >50 mm).')
add_bullet('XAI explanation: Plain-English explanation of why the model made each prediction, based on SHAP analysis.')
add_bullet('SHAP chart: Visual feature contribution chart for each prediction.')
add_para('To run: pip install -r requirements.txt && streamlit run app.py')

# ── 7. References ──
add_heading('7. References', 1)
refs = [
    'Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785\u2013794.',
    'Lundberg, S. M., & Lee, S.-I. (2017). A Unified Approach to Interpreting Model Predictions. Advances in Neural Information Processing Systems, 30, 4765\u20134774.',
    'Open-Meteo. (2023). Historical Weather API. https://open-meteo.com/en/docs/historical-weather-api',
    'Department of Meteorology, Sri Lanka. Climatological data and monsoon patterns. http://www.meteo.gov.lk/',
    'Akiba, T., et al. (2019). Optuna: A Next-generation Hyperparameter Optimization Framework. Proceedings of the 25th ACM SIGKDD, 2623\u20132631.',
    'Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. JMLR, 12, 2825\u20132830.',
]
for i, ref in enumerate(refs, 1):
    add_para(f'[{i}] {ref}')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Project developed as part of a Machine Learning assignment. '
               'All data used is publicly available and contains no personal or sensitive information.')
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Save
doc.save('Documentation.docx')
print("\nDocumentation.docx created successfully!")
print(f"Images saved in: {IMG_DIR}/")
